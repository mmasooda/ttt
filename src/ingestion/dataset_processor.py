import fitz  # PyMuPDF
import pandas as pd
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
import json
import re
from dataclasses import dataclass, asdict
import hashlib
from PIL import Image
import docx2txt
import io

from ..utils import logger, S3Client

@dataclass
class ProjectDocument:
    """Single project document"""
    filename: str
    file_type: str  # 'specs', 'boq', 'offer', 'other'
    content: str
    metadata: Dict[str, Any]
    entities: List[Dict[str, Any]]
    tables: List[Dict[str, Any]]
    images: List[Dict[str, Any]]

@dataclass
class ProjectSet:
    """Complete project set with specs, BOQ, and offer"""
    project_id: str
    project_name: str
    specs_doc: Optional[ProjectDocument]
    boq_doc: Optional[ProjectDocument]
    offer_doc: Optional[ProjectDocument]
    other_docs: List[ProjectDocument]
    summary: Dict[str, Any]

class DatasetProcessor:
    """Process dataset directory for TTT inference samples"""
    
    def __init__(self):
        self.s3_client = S3Client()
        self.supported_extensions = {
            '.pdf': self._process_pdf,
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
            '.docx': self._process_word,
            '.doc': self._process_word,
            '.png': self._process_image,
            '.jpg': self._process_image,
            '.jpeg': self._process_image
        }
        
        # Keywords to identify document types
        self.doc_type_keywords = {
            'specs': ['spec', 'specification', 'requirement', 'standard', 'compliance'],
            'boq': ['boq', 'bill', 'quantity', 'quantities', 'pricing', 'cost'],
            'offer': ['offer', 'proposal', 'quotation', 'quote', 'price']
        }
    
    def discover_dataset_structure(self, dataset_prefix: str = "dataset/") -> Dict[str, List[Dict[str, Any]]]:
        """Discover and organize dataset files by project"""
        try:
            logger.info("Discovering dataset structure", prefix=dataset_prefix)
            
            # Get all files in dataset directory
            all_objects = self.s3_client.list_all_objects(prefix=dataset_prefix)
            
            # Group files by directory/project
            projects = {}
            
            for obj in all_objects:
                key = obj['key']
                
                # Skip if not supported file type
                ext = Path(key).suffix.lower()
                if ext not in self.supported_extensions:
                    continue
                
                # Extract project/directory info
                path_parts = key.replace(dataset_prefix, '').split('/')
                
                if len(path_parts) > 1:
                    # File is in a subdirectory
                    project_dir = path_parts[0]
                    if project_dir not in projects:
                        projects[project_dir] = []
                    projects[project_dir].append(obj)
                else:
                    # File is in root dataset directory
                    if 'root' not in projects:
                        projects['root'] = []
                    projects['root'].append(obj)
            
            logger.info("Dataset structure discovered", projects=len(projects))
            
            # Print structure
            for project, files in projects.items():
                logger.info(f"Project: {project}", files=len(files))
                for file_obj in files:
                    logger.info(f"  - {file_obj['key']} ({file_obj['size']/1024/1024:.2f} MB)")
            
            return projects
            
        except Exception as e:
            logger.error("Failed to discover dataset structure", error=str(e))
            return {}
    
    def identify_document_type(self, filename: str, content: str = "") -> str:
        """Identify document type based on filename and content"""
        filename_lower = filename.lower()
        content_lower = content.lower()
        
        scores = {}
        for doc_type, keywords in self.doc_type_keywords.items():
            score = 0
            for keyword in keywords:
                if keyword in filename_lower:
                    score += 3  # Filename match is more important
                if keyword in content_lower:
                    score += 1
            scores[doc_type] = score
        
        # Return type with highest score, or 'other' if no clear match
        if scores:
            best_type = max(scores, key=scores.get)
            if scores[best_type] > 0:
                return best_type
        
        return 'other'
    
    async def process_project_files(self, project_name: str, file_objects: List[Dict[str, Any]]) -> ProjectSet:
        """Process all files for a single project"""
        try:
            logger.info("Processing project files", project=project_name, files=len(file_objects))
            
            # Download and process each file
            documents = []
            download_dir = Path(f"./data/downloads/dataset/{project_name}")
            
            download_results = await self.s3_client.download_files_batch(
                file_objects, download_dir, max_concurrent=3
            )
            
            for file_result in download_results:
                if not file_result['download_success']:
                    continue
                
                local_path = Path(file_result['local_path'])
                doc = await self._process_single_document(local_path, file_result['key'])
                
                if doc:
                    documents.append(doc)
            
            # Organize documents by type
            specs_doc = None
            boq_doc = None
            offer_doc = None
            other_docs = []
            
            for doc in documents:
                if doc.file_type == 'specs':
                    specs_doc = doc
                elif doc.file_type == 'boq':
                    boq_doc = doc
                elif doc.file_type == 'offer':
                    offer_doc = doc
                else:
                    other_docs.append(doc)
            
            # Create project summary
            summary = self._create_project_summary(documents)
            
            project_set = ProjectSet(
                project_id=hashlib.md5(project_name.encode()).hexdigest()[:8],
                project_name=project_name,
                specs_doc=specs_doc,
                boq_doc=boq_doc,
                offer_doc=offer_doc,
                other_docs=other_docs,
                summary=summary
            )
            
            logger.info("Project processed", 
                       project=project_name,
                       specs=specs_doc is not None,
                       boq=boq_doc is not None,
                       offer=offer_doc is not None,
                       other=len(other_docs))
            
            return project_set
            
        except Exception as e:
            logger.error("Failed to process project", project=project_name, error=str(e))
            return None
    
    async def _process_single_document(self, file_path: Path, s3_key: str) -> Optional[ProjectDocument]:
        """Process a single document file"""
        try:
            ext = file_path.suffix.lower()
            if ext not in self.supported_extensions:
                return None
            
            processor = self.supported_extensions[ext]
            content, metadata, tables, images = await processor(file_path)
            
            # Identify document type
            doc_type = self.identify_document_type(file_path.name, content)
            
            # Extract entities
            entities = self._extract_entities(content, doc_type)
            
            return ProjectDocument(
                filename=file_path.name,
                file_type=doc_type,
                content=content,
                metadata=metadata,
                entities=entities,
                tables=tables,
                images=images
            )
            
        except Exception as e:
            logger.error("Failed to process document", file=str(file_path), error=str(e))
            return None
    
    async def _process_pdf(self, file_path: Path) -> Tuple[str, Dict, List, List]:
        """Process PDF file"""
        try:
            doc = fitz.open(str(file_path))
            content_parts = []
            tables = []
            images = []
            
            metadata = {
                'page_count': doc.page_count,
                'title': doc.metadata.get('title', ''),
                'author': doc.metadata.get('author', ''),
                'subject': doc.metadata.get('subject', '')
            }
            
            for page_num in range(doc.page_count):
                page = doc.page(page_num)
                
                # Extract text
                text = page.get_text()
                if text.strip():
                    content_parts.append(f"[Page {page_num + 1}]\n{text}")
                
                # Extract tables
                try:
                    page_tables = page.find_tables()
                    for table in page_tables:
                        table_data = table.extract()
                        if table_data:
                            tables.append({
                                'page': page_num + 1,
                                'data': table_data,
                                'bbox': table.bbox
                            })
                except:
                    pass  # Table extraction may fail on some PDFs
                
                # Extract images
                try:
                    image_list = page.get_images()
                    for img_index, img in enumerate(image_list):
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        if pix.n < 5:  # GRAY or RGB
                            images.append({
                                'page': page_num + 1,
                                'index': img_index,
                                'width': pix.width,
                                'height': pix.height,
                                'colorspace': pix.colorspace.name if pix.colorspace else 'unknown'
                            })
                        pix = None
                except:
                    pass
            
            doc.close()
            content = "\n\n".join(content_parts)
            
            return content, metadata, tables, images
            
        except Exception as e:
            logger.error("PDF processing failed", file=str(file_path), error=str(e))
            return "", {}, [], []
    
    async def _process_excel(self, file_path: Path) -> Tuple[str, Dict, List, List]:
        """Process Excel file"""
        try:
            excel_file = pd.ExcelFile(file_path)
            content_parts = []
            tables = []
            
            metadata = {
                'sheet_count': len(excel_file.sheet_names),
                'sheet_names': excel_file.sheet_names
            }
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Content for text analysis
                content_parts.append(f"=== Sheet: {sheet_name} ===")
                content_parts.append(f"Columns: {', '.join(df.columns.tolist())}")
                content_parts.append(df.to_string(index=False, max_rows=20))
                
                # Structured table data
                tables.append({
                    'sheet_name': sheet_name,
                    'columns': df.columns.tolist(),
                    'data': df.head(50).to_dict('records'),  # Limit for JSON size
                    'row_count': len(df),
                    'summary_stats': df.describe().to_dict() if df.select_dtypes(include='number').shape[1] > 0 else {}
                })
            
            content = "\n\n".join(content_parts)
            return content, metadata, tables, []
            
        except Exception as e:
            logger.error("Excel processing failed", file=str(file_path), error=str(e))
            return "", {}, [], []
    
    async def _process_word(self, file_path: Path) -> Tuple[str, Dict, List, List]:
        """Process Word document"""
        try:
            # Extract text content
            content = docx2txt.process(str(file_path))
            
            metadata = {
                'word_count': len(content.split()),
                'char_count': len(content)
            }
            
            # Try to extract tables (basic)
            tables = []
            try:
                from docx import Document
                doc = Document(str(file_path))
                
                for i, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)
                    
                    if table_data:
                        tables.append({
                            'table_index': i,
                            'data': table_data,
                            'row_count': len(table_data),
                            'col_count': len(table_data[0]) if table_data else 0
                        })
                        
                metadata['table_count'] = len(tables)
                
            except ImportError:
                logger.warning("python-docx not available, skipping table extraction")
            except Exception as e:
                logger.warning("Word table extraction failed", error=str(e))
            
            return content, metadata, tables, []
            
        except Exception as e:
            logger.error("Word processing failed", file=str(file_path), error=str(e))
            return "", {}, [], []
    
    async def _process_image(self, file_path: Path) -> Tuple[str, Dict, List, List]:
        """Process image file"""
        try:
            with Image.open(file_path) as img:
                metadata = {
                    'width': img.width,
                    'height': img.height,
                    'format': img.format,
                    'mode': img.mode
                }
                
                # For images, content is description
                content = f"Image: {file_path.name} ({img.width}x{img.height}, {img.format})"
                
                images = [{
                    'filename': file_path.name,
                    'width': img.width,
                    'height': img.height,
                    'format': img.format,
                    'mode': img.mode
                }]
                
                return content, metadata, [], images
                
        except Exception as e:
            logger.error("Image processing failed", file=str(file_path), error=str(e))
            return "", {}, [], []
    
    def _extract_entities(self, content: str, doc_type: str) -> List[Dict[str, Any]]:
        """Extract entities specific to document type"""
        entities = []
        
        # Common patterns
        patterns = {
            'product_code': r'\b[A-Z]{1,4}[-_]?\d{2,6}[A-Z]?\b',
            'money': r'\$?[\d,]+\.?\d*\s*(USD|SAR|AED|QAR)?',
            'percentage': r'\d+\.?\d*%',
            'quantity': r'\b\d+\s*(pcs?|units?|nos?|sets?|each)\b',
            'area': r'\d+\.?\d*\s*(m2|sqm|square\s*meters?)',
            'voltage': r'\d+\s*V\b',
            'current': r'\d+\.?\d*\s*A\b'
        }
        
        # Document-specific patterns
        if doc_type == 'specs':
            patterns.update({
                'standard': r'\b(NFPA|BS|EN|IEC|ANSI)\s*\d+',
                'compliance': r'\b(compliant|compliance|conform|accordance)\b',
                'requirement': r'\b(shall|must|required|mandatory)\b'
            })
        elif doc_type == 'boq':
            patterns.update({
                'item_description': r'Item\s*\d+:?\s*([^\n]+)',
                'total_cost': r'(total|grand\s*total|sum):\s*[\$\d,]+',
                'unit_price': r'unit\s*price:\s*[\$\d,]+'
            })
        
        for entity_type, pattern in patterns.items():
            matches = re.finditer(pattern, content, re.IGNORECASE)
            for match in matches:
                entities.append({
                    'name': match.group().strip(),
                    'type': entity_type,
                    'confidence': 0.8,
                    'start_pos': match.start(),
                    'end_pos': match.end(),
                    'context': content[max(0, match.start()-50):match.end()+50]
                })
        
        return entities
    
    def _create_project_summary(self, documents: List[ProjectDocument]) -> Dict[str, Any]:
        """Create project summary from documents"""
        summary = {
            'total_documents': len(documents),
            'document_types': {},
            'total_entities': 0,
            'total_tables': 0,
            'total_images': 0,
            'key_products': [],
            'estimated_value': None,
            'standards_mentioned': []
        }
        
        all_entities = []
        
        for doc in documents:
            doc_type = doc.file_type
            if doc_type not in summary['document_types']:
                summary['document_types'][doc_type] = 0
            summary['document_types'][doc_type] += 1
            
            summary['total_entities'] += len(doc.entities)
            summary['total_tables'] += len(doc.tables)
            summary['total_images'] += len(doc.images)
            
            all_entities.extend(doc.entities)
        
        # Extract key information
        product_codes = [e['name'] for e in all_entities if e['type'] == 'product_code']
        money_values = [e['name'] for e in all_entities if e['type'] == 'money']
        standards = [e['name'] for e in all_entities if e['type'] == 'standard']
        
        summary['key_products'] = list(set(product_codes))[:10]  # Top 10 unique products
        summary['estimated_value'] = money_values[0] if money_values else None
        summary['standards_mentioned'] = list(set(standards))
        
        return summary
    
    def export_to_json(self, project_sets: List[ProjectSet], output_file: Path) -> bool:
        """Export processed project sets to JSON for TTT inference"""
        try:
            # Convert to serializable format
            json_data = {
                'dataset_info': {
                    'total_projects': len(project_sets),
                    'processed_at': pd.Timestamp.now().isoformat(),
                    'version': '1.0'
                },
                'projects': []
            }
            
            for project in project_sets:
                project_dict = asdict(project)
                json_data['projects'].append(project_dict)
            
            # Write to file
            output_file.parent.mkdir(parents=True, exist_ok=True)
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(json_data, f, indent=2, ensure_ascii=False, default=str)
            
            logger.info("Dataset exported to JSON", 
                       file=str(output_file), 
                       projects=len(project_sets))
            
            return True
            
        except Exception as e:
            logger.error("Failed to export to JSON", error=str(e))
            return False